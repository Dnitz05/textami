#!/usr/bin/env python3
"""
TEXTAMI OOXML Parser - ingest_docx.py
=====================================

Parseja documents DOCX per extreure:
1. styles.xml - Definicions d'estils Word
2. numbering.xml - Definicions de llistes i numeració  
3. document.xml - Contingut i estructura
4. Genera styleManifest.json amb mappings HTML semàntic

Author: OOXML-Centric Architecture
Date: 30 Agost 2025
"""

import sys
import json
import zipfile
import os
from pathlib import Path
from typing import Dict, List, Optional, Any, Tuple
import argparse
from datetime import datetime

try:
    from lxml import etree
    import docx
    from docx import Document
except ImportError as e:
    print(f" Error: Missing required dependencies: {e}")
    print("Install with: pip install python-docx lxml")
    sys.exit(1)

# Namespaces OOXML
NAMESPACES = {
    'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
    'w14': 'http://schemas.microsoft.com/office/word/2010/wordml',
    'mc': 'http://schemas.openxmlformats.org/markup-compatibility/2006'
}

class OOXMLStyleParser:
    """Parser per extreure i processar estils OOXML"""
    
    def __init__(self, docx_path: str):
        """
        Inicialitza el parser amb el path al DOCX
        
        Args:
            docx_path: Path al fitxer DOCX
        """
        self.docx_path = docx_path
        self.styles_data = {}
        self.numbering_data = {}
        self.document_structure = []
        self.style_manifest = {}
        
        # Vocabulari HTML semàntic estàndard
        self.html_vocabulary = {
            'h1': 'Heading 1, Title, Heading1Char',
            'h2': 'Heading 2, Subtitle, Heading2Char', 
            'h3': 'Heading 3, Heading3Char',
            'p': 'Normal, BodyText, DefaultParagraphFont',
            'p.BodyText': 'Body Text, BodyTextChar',
            'ul.Bulleted': 'List Bullet, ListBullet, BulletList',
            'ol.Numbered': 'List Number, ListNumber, NumberedList',
            'table.StdTable': 'Table Grid, TableGrid, StandardTable',
            'blockquote': 'Quote, BlockText, IntenseQuote'
        }
        
        print(f"OOXML Parser initialized for: {docx_path}")
    
    def extract_styles_xml(self) -> Dict[str, Any]:
        """
        Extreu i parseja styles.xml del DOCX
        
        Returns:
            Dictionary amb estils processats
        """
        print("📝 Extracting styles.xml...")
        
        try:
            with zipfile.ZipFile(self.docx_path, 'r') as docx_zip:
                # Llegir styles.xml
                styles_xml = docx_zip.read('word/styles.xml')
                styles_root = etree.fromstring(styles_xml)
                
                styles = {}
                
                # Processar cada estil
                for style_elem in styles_root.xpath('//w:style', namespaces=NAMESPACES):
                    style_id = style_elem.get('{%s}styleId' % NAMESPACES['w'])
                    style_type = style_elem.get('{%s}type' % NAMESPACES['w'])
                    
                    # Obtenir nom de l'estil
                    name_elem = style_elem.xpath('./w:name', namespaces=NAMESPACES)
                    style_name = name_elem[0].get('{%s}val' % NAMESPACES['w']) if name_elem else style_id
                    
                    # Obtenir basedOn (herència d'estils)
                    based_on_elem = style_elem.xpath('./w:basedOn', namespaces=NAMESPACES)
                    based_on = based_on_elem[0].get('{%s}val' % NAMESPACES['w']) if based_on_elem else None
                    
                    styles[style_id] = {
                        'name': style_name,
                        'type': style_type,
                        'basedOn': based_on,
                        'xmlElement': style_elem
                    }
                
                self.styles_data = styles
                print(f" Extracted {len(styles)} styles from styles.xml")
                return styles
                
        except Exception as e:
            print(f" Error extracting styles.xml: {e}")
            return {}
    
    def extract_numbering_xml(self) -> Dict[str, Any]:
        """
        Extreu i parseja numbering.xml per llistes
        
        Returns:
            Dictionary amb definicions de numbering
        """
        print("📝 Extracting numbering.xml...")
        
        try:
            with zipfile.ZipFile(self.docx_path, 'r') as docx_zip:
                # Comprovar si existeix numbering.xml
                if 'word/numbering.xml' not in docx_zip.namelist():
                    print(" No numbering.xml found - document has no lists")
                    return {}
                
                numbering_xml = docx_zip.read('word/numbering.xml')
                numbering_root = etree.fromstring(numbering_xml)
                
                numbering = {}
                
                # Processar definicions de numbering
                for num_elem in numbering_root.xpath('//w:num', namespaces=NAMESPACES):
                    num_id = num_elem.get('{%s}numId' % NAMESPACES['w'])
                    
                    # Obtenir abstractNum referència
                    abstract_num_elem = num_elem.xpath('./w:abstractNumId', namespaces=NAMESPACES)
                    abstract_num_id = abstract_num_elem[0].get('{%s}val' % NAMESPACES['w']) if abstract_num_elem else None
                    
                    numbering[num_id] = {
                        'abstractNumId': abstract_num_id,
                        'xmlElement': num_elem
                    }
                
                self.numbering_data = numbering
                print(f" Extracted {len(numbering)} numbering definitions")
                return numbering
                
        except Exception as e:
            print(f" Error extracting numbering.xml: {e}")
            return {}
    
    def analyze_document_structure(self) -> List[Dict[str, Any]]:
        """
        Analitza l'estructura del document principal
        
        Returns:
            Llista d'elements del document amb estils
        """
        print("📄 Analyzing document structure...")
        
        try:
            # Usar python-docx per llegir contingut
            doc = Document(self.docx_path)
            structure = []
            
            for paragraph in doc.paragraphs:
                # Obtenir estil del paràgraf
                style_name = paragraph.style.name if paragraph.style else 'Normal'
                
                element = {
                    'type': 'paragraph',
                    'text': paragraph.text[:100] + '...' if len(paragraph.text) > 100 else paragraph.text,
                    'style': style_name,
                    'full_text': paragraph.text
                }
                structure.append(element)
            
            # Processar taules
            for table in doc.tables:
                table_structure = {
                    'type': 'table',
                    'rows': len(table.rows),
                    'cols': len(table.columns) if table.rows else 0,
                    'style': 'TableGrid'  # Default table style
                }
                structure.append(table_structure)
            
            self.document_structure = structure
            print(f" Analyzed document: {len(structure)} elements")
            return structure
            
        except Exception as e:
            print(f" Error analyzing document structure: {e}")
            return []
    
    def generate_html_semantic_mappings(self) -> Dict[str, Any]:
        """
        Genera mappings heurístics d'estils Word → HTML semàntic
        
        Returns:
            Dictionary amb mappings HTML
        """
        print("🎯 Generating HTML semantic mappings...")
        
        mappings = {}
        warnings = []
        fallbacks = {}
        
        for style_id, style_data in self.styles_data.items():
            style_name = style_data['name']
            style_type = style_data['type']
            
            # Skip character styles i altres no rellevants
            if style_type not in ['paragraph', 'table']:
                continue
            
            # Heurística per mapping
            html_element = self._map_style_to_html(style_name, style_type)
            
            if html_element:
                mappings[html_element] = style_name
            else:
                # Fallback per estils no reconeguts
                fallback_element = self._generate_fallback_mapping(style_name, style_type)
                fallbacks[style_name] = fallback_element
                warnings.append(f"Unknown style '{style_name}' mapped to '{fallback_element}'")
        
        # Assegurar que tenim mappings bàsics
        self._ensure_basic_mappings(mappings)
        
        manifest = {
            'version': '1.0',
            'generatedAt': datetime.now().isoformat(),
            'docx_path': os.path.basename(self.docx_path),
            'styles': mappings,
            'fallbacks': fallbacks,
            'warnings': warnings,
            'vocabulary': list(self.html_vocabulary.keys()),
            'statistics': {
                'total_styles_found': len(self.styles_data),
                'mapped_styles': len(mappings),
                'fallback_styles': len(fallbacks)
            }
        }
        
        self.style_manifest = manifest
        print(f" Generated {len(mappings)} HTML mappings with {len(fallbacks)} fallbacks")
        return manifest
    
    def _map_style_to_html(self, style_name: str, style_type: str) -> Optional[str]:
        """
        Mapeja un estil Word a element HTML semàntic
        
        Args:
            style_name: Nom de l'estil Word
            style_type: Tipus d'estil (paragraph, table, etc.)
            
        Returns:
            Element HTML semàntic o None
        """
        style_lower = style_name.lower()
        
        # Mapings explícits per headings
        if 'heading 1' in style_lower or style_lower == 'title':
            return 'h1'
        elif 'heading 2' in style_lower or 'subtitle' in style_lower:
            return 'h2'  
        elif 'heading 3' in style_lower:
            return 'h3'
        
        # Mapings per text body
        elif 'body' in style_lower or 'normal' in style_lower:
            return 'p'
        elif 'body text' in style_lower:
            return 'p.BodyText'
        
        # Mapings per llistes
        elif 'list bullet' in style_lower or 'bullet' in style_lower:
            return 'ul.Bulleted'
        elif 'list number' in style_lower or 'numbered' in style_lower:
            return 'ol.Numbered'
        
        # Mapings per taules
        elif style_type == 'table' or 'table' in style_lower:
            return 'table.StdTable'
        
        # Mapings per quotes
        elif 'quote' in style_lower or 'block' in style_lower:
            return 'blockquote'
        
        return None
    
    def _generate_fallback_mapping(self, style_name: str, style_type: str) -> str:
        """
        Genera mapping de fallback per estils no reconeguts
        
        Args:
            style_name: Nom de l'estil
            style_type: Tipus d'estil
            
        Returns:
            Element HTML de fallback
        """
        if style_type == 'table':
            return 'table.StdTable'
        else:
            return 'p'  # Fallback per defecte
    
    def _ensure_basic_mappings(self, mappings: Dict[str, str]) -> None:
        """
        Assegura que existeixen els mappings bàsics essencials
        
        Args:
            mappings: Dictionary de mappings a modificar
        """
        essentials = {
            'h1': 'Heading 1',
            'h2': 'Heading 2', 
            'p': 'Normal',
            'table.StdTable': 'Table Grid'
        }
        
        for html_element, default_style in essentials.items():
            if html_element not in mappings:
                mappings[html_element] = default_style
    
    def generate_html_preview(self) -> str:
        """
        Genera preview HTML del document amb estils semàntics
        
        Returns:
            HTML string del document
        """
        print("🌐 Generating HTML preview...")
        
        html_parts = ['<!DOCTYPE html>', '<html>', '<head>']
        html_parts.append('<meta charset="UTF-8">')
        html_parts.append('<title>DOCX Preview - Textami</title>')
        html_parts.append('<style>')
        html_parts.append(self._generate_preview_css())
        html_parts.append('</style>')
        html_parts.extend(['</head>', '<body>'])
        html_parts.append('<div class="document-preview">')
        
        # Processar elements del document
        for element in self.document_structure:
            if element['type'] == 'paragraph' and element['text'].strip():
                html_element = self._style_to_html_element(element['style'])
                text = element['text']
                html_parts.append(f'<{html_element}>{text}</{html_element}>')
            elif element['type'] == 'table':
                html_parts.append(f'<table class="StdTable"><tr><td>[Table {element["rows"]}x{element["cols"]}]</td></tr></table>')
        
        html_parts.extend(['</div>', '</body>', '</html>'])
        
        html_content = '\n'.join(html_parts)
        print(" Generated HTML preview")
        return html_content
    
    def _style_to_html_element(self, style_name: str) -> str:
        """
        Converteix estil Word a element HTML
        
        Args:
            style_name: Nom de l'estil Word
            
        Returns:
            Element HTML corresponent
        """
        # Buscar en els mappings generats
        for html_elem, mapped_style in self.style_manifest.get('styles', {}).items():
            if mapped_style == style_name:
                return html_elem.split('.')[0]  # Treure classes CSS
        
        # Fallback
        return 'p'
    
    def _generate_preview_css(self) -> str:
        """
        Genera CSS bàsic per preview HTML
        
        Returns:
            CSS string
        """
        return """
        .document-preview {
            max-width: 800px;
            margin: 20px auto;
            padding: 20px;
            font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
            line-height: 1.6;
        }
        h1 { color: #2c3e50; border-bottom: 2px solid #3498db; }
        h2 { color: #34495e; border-bottom: 1px solid #bdc3c7; }
        h3 { color: #34495e; }
        .BodyText { margin: 10px 0; }
        .StdTable { border-collapse: collapse; width: 100%; margin: 10px 0; }
        .StdTable td { border: 1px solid #ddd; padding: 8px; }
        blockquote { border-left: 4px solid #3498db; margin: 10px 0; padding-left: 15px; }
        """
    
    def save_outputs(self, output_dir: str) -> Dict[str, str]:
        """
        Desa tots els outputs generats
        
        Args:
            output_dir: Directory on desar els fitxers
            
        Returns:
            Dictionary amb paths dels fitxers generats
        """
        print(f"💾 Saving outputs to {output_dir}...")
        
        # Crear directory si no existeix
        Path(output_dir).mkdir(parents=True, exist_ok=True)
        
        base_name = Path(self.docx_path).stem
        outputs = {}
        
        # Desar styleManifest.json
        manifest_path = os.path.join(output_dir, f"{base_name}_styleManifest.json")
        with open(manifest_path, 'w', encoding='utf-8') as f:
            json.dump(self.style_manifest, f, indent=2, ensure_ascii=False)
        outputs['manifest'] = manifest_path
        
        # Desar HTML preview
        html_preview = self.generate_html_preview()
        html_path = os.path.join(output_dir, f"{base_name}_preview.html")
        with open(html_path, 'w', encoding='utf-8') as f:
            f.write(html_preview)
        outputs['html_preview'] = html_path
        
        # Desar processing report
        report = {
            'processing_date': datetime.now().isoformat(),
            'docx_file': self.docx_path,
            'styles_found': len(self.styles_data),
            'numbering_found': len(self.numbering_data),
            'document_elements': len(self.document_structure),
            'html_mappings': len(self.style_manifest.get('styles', {})),
            'warnings': self.style_manifest.get('warnings', []),
            'success': True
        }
        
        report_path = os.path.join(output_dir, f"{base_name}_report.json")
        with open(report_path, 'w', encoding='utf-8') as f:
            json.dump(report, f, indent=2, ensure_ascii=False)
        outputs['report'] = report_path
        
        print(f" Saved outputs: {', '.join(outputs.keys())}")
        return outputs
    
    def process_full_pipeline(self, output_dir: Optional[str] = None) -> Dict[str, Any]:
        """
        Executa el pipeline complet d'OOXML processing
        
        Args:
            output_dir: Directory opcional per outputs
            
        Returns:
            Dictionary amb tots els resultats
        """
        print(f"Starting OOXML processing pipeline for: {self.docx_path}")
        start_time = datetime.now()
        
        try:
            # 1. Extreure styles.xml
            self.extract_styles_xml()
            
            # 2. Extreure numbering.xml
            self.extract_numbering_xml()
            
            # 3. Analitzar estructura document
            self.analyze_document_structure()
            
            # 4. Generar mappings HTML
            self.generate_html_semantic_mappings()
            
            # 5. Desar outputs si requerit
            outputs = {}
            if output_dir:
                outputs = self.save_outputs(output_dir)
            
            processing_time = (datetime.now() - start_time).total_seconds()
            
            result = {
                'success': True,
                'processing_time_seconds': processing_time,
                'style_manifest': self.style_manifest,
                'document_structure': self.document_structure,
                'outputs': outputs
            }
            
            print(f" OOXML processing completed in {processing_time:.2f} seconds")
            return result
            
        except Exception as e:
            error_result = {
                'success': False,
                'error': str(e),
                'processing_time_seconds': (datetime.now() - start_time).total_seconds()
            }
            print(f" OOXML processing failed: {e}")
            return error_result

def main():
    """Entry point per command line usage"""
    parser = argparse.ArgumentParser(description='Textami OOXML Parser')
    parser.add_argument('docx_file', help='Path to DOCX file to process')
    parser.add_argument('--output-dir', '-o', help='Output directory for generated files')
    parser.add_argument('--verbose', '-v', action='store_true', help='Verbose output')
    
    args = parser.parse_args()
    
    if not os.path.exists(args.docx_file):
        print(f" Error: File not found: {args.docx_file}")
        sys.exit(1)
    
    # Crear parser i processar
    parser = OOXMLStyleParser(args.docx_file)
    result = parser.process_full_pipeline(args.output_dir)
    
    if result['success']:
        print(f"Success! Processing time: {result['processing_time_seconds']:.2f}s")
        if args.verbose:
            print(json.dumps(result['style_manifest'], indent=2))
        sys.exit(0)
    else:
        print(f"Failed: {result['error']}")
        sys.exit(1)

if __name__ == '__main__':
    main()
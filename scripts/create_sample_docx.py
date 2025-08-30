#!/usr/bin/env python3
"""
Create a sample DOCX file for testing OOXML parser
"""

from docx import Document
from docx.shared import Inches

def create_sample_docx():
    """Create a sample DOCX with various styles for testing"""
    
    # Create new document
    doc = Document()
    
    # Add title
    title = doc.add_heading('Document de Prova Textami', 0)
    
    # Add subtitle
    doc.add_heading('Subtítol del Document', level=2)
    
    # Add normal paragraph
    doc.add_paragraph('Aquest és un paràgraf normal amb text de prova.')
    
    # Add body text paragraph
    p = doc.add_paragraph('Aquest és text del cos del document amb estil específic.')
    
    # Add bullet list
    doc.add_paragraph('Element de llista 1', style='List Bullet')
    doc.add_paragraph('Element de llista 2', style='List Bullet')
    doc.add_paragraph('Element de llista 3', style='List Bullet')
    
    # Add numbered list
    doc.add_paragraph('Primer element numerat', style='List Number')
    doc.add_paragraph('Segon element numerat', style='List Number')
    doc.add_paragraph('Tercer element numerat', style='List Number')
    
    # Add another heading
    doc.add_heading('Secció amb Taula', level=3)
    
    # Add table
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Nom'
    hdr_cells[1].text = 'Cognoms'
    hdr_cells[2].text = 'Edat'
    
    # Add data rows
    row_data = [
        ['Joan', 'García', '25'],
        ['Maria', 'López', '30'],
        ['Pere', 'Martí', '35']
    ]
    
    for name, surname, age in row_data:
        row_cells = table.add_row().cells
        row_cells[0].text = name
        row_cells[1].text = surname
        row_cells[2].text = age
    
    # Add quote
    doc.add_paragraph(
        'Aquesta és una cita o text destacat que hauria de ser reconegut com a blockquote.',
        style='Quote'
    )
    
    # Add final paragraph
    doc.add_paragraph('Paràgraf final del document de prova.')
    
    # Save document
    doc.save('sample_document.docx')
    print("SUCCESS: Created sample_document.docx")
    
    # Print summary
    print("Sample document contains:")
    print("- 1 main title (Heading 1)")
    print("- 1 subtitle (Heading 2)")  
    print("- 1 section heading (Heading 3)")
    print("- Multiple normal paragraphs")
    print("- Bullet list (3 items)")
    print("- Numbered list (3 items)")
    print("- Table with headers (3x4)")
    print("- Quote paragraph")

if __name__ == '__main__':
    create_sample_docx()